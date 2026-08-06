"""Build the teammate-facing Word user guide."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Apple SD Gothic Neo"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "263238"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GOLD = "FFF4D6"


def _font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)


def _cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        if node.getparent() is None:
            tc_mar.append(node)


def _set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")
    if tc_w.getparent() is None:
        tc_pr.append(tc_w)


def _table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    if tbl_ind.getparent() is None:
        tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[i])
            _cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_text(cell, text, bold=False, color=INK):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _font(run, 10.2, bold, color)


def _heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def _body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        _font(r, 11, True, INK)
        text = text[len(bold_lead):]
    r = p.add_run(text)
    _font(r, 11, False, INK)
    return p


def _new_numbering(doc, num_fmt, marker, left=540, hanging=270):
    root = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in root.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), marker)
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    root.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    root.append(num)
    return num_id


def _apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def _bullet(doc, text, num_id):
    p = doc.add_paragraph()
    _apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    _font(r, 11, False, INK)
    return p


def _number(doc, lead, text, num_id):
    p = doc.add_paragraph()
    _apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(lead)
    _font(r, 11, True, INK)
    r = p.add_run(text)
    _font(r, 11, False, INK)
    return p


def _callout(doc, label, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    _table_geometry(table, [9360])
    cell = table.cell(0, 0)
    _shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(label + "  ")
    _font(r, 10.5, True, DARK_BLUE)
    r = p.add_run(text)
    _font(r, 10.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("계약서 리뷰 가이드  |  ")
    _font(run, 9, False, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def build_guide(version, out_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    _add_page_number(sec.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("계약서 리뷰 가이드")
    _font(r, 25, True, DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("팀원용 사용안내  |  v" + version)
    _font(r, 11, False, MUTED)

    _callout(doc, "한 줄 설명", "계약서에서 검토할 항목과 관련 원문을 찾아주는 보조 도구입니다. 최종 결론은 검토자가 직접 남깁니다.", LIGHT_BLUE)

    _heading(doc, "가장 빠른 사용 순서", 1)
    quick_num = _new_numbering(doc, "decimal", "%1.")
    steps = [
        ("파일 열기", "계약서 파일을 열거나 본문을 붙여넣습니다."),
        ("설정 확인", "당사 지위, 계약 유형, 적용 모듈을 확인합니다."),
        ("분석 시작", "조항별 검토 화면에서 관련 원문과 검토 항목을 확인합니다."),
        ("판정 입력", "각 항목을 이상없음 또는 검토의견으로 판정하고 코멘트를 남깁니다."),
        ("결과 공유", "종합리포트를 확인한 뒤 검토의견 파일로 내보냅니다."),
    ]
    for lead, text in steps:
        _number(doc, lead + "  ", text, quick_num)

    _callout(doc, "중요", "입력한 계약서는 외부로 전송되지 않습니다. 이 도구는 검토를 돕지만 법적 판단을 대신하지 않습니다.", GOLD)

    _heading(doc, "1. 시작하기", 1)
    start_num = _new_numbering(doc, "decimal", "%1.")
    _number(doc, "압축 풀기  ", "배포받은 ZIP 파일의 압축을 풉니다.", start_num)
    _number(doc, "앱 열기  ", "contract-review.html을 Chrome 또는 Edge로 엽니다.", start_num)
    _number(doc, "계약서 넣기  ", "[파일 열기]를 누르거나 계약서 내용을 입력 칸에 붙여넣습니다.", start_num)
    _number(doc, "설정 확인  ", "당사 지위와 계약 유형이 맞는지 확인하고 필요한 적용 모듈을 선택합니다.", start_num)
    _number(doc, "분석 실행  ", "[분석 시작]을 누릅니다.", start_num)

    _heading(doc, "지원 파일", 2)
    _body(doc, "PDF, DOCX, DOC, HWP, HWPX 파일을 열 수 있습니다.")
    _callout(doc, "파일을 읽지 못할 때", "스캔 PDF, 암호 문서 또는 표가 복잡한 문서는 글자를 제대로 읽지 못할 수 있습니다. 이 경우 문서의 텍스트를 복사해 입력 칸에 붙여넣으세요.")

    _heading(doc, "2. 분석 전에 확인할 설정", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = ("설정", "확인할 내용")
    for i, text in enumerate(headers):
        _shade(table.rows[0].cells[i], LIGHT_BLUE)
        _add_text(table.rows[0].cells[i], text, True, DARK_BLUE)
    rows = [
        ("당사 지위", "당사가 의무를 부담하는 계약 당사자인지, 수익자·투자자인지 확인합니다."),
        ("계약 유형", "업무위탁, 투자, 비밀유지 등 실제 계약 성격과 맞는지 확인합니다."),
        ("적용 모듈", "개인정보, 전자금융, 하도급 등 이 계약에 필요한 영역만 선택합니다."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        _add_text(cells[0], left, True)
        _add_text(cells[1], right)
    _table_geometry(table, [2700, 6660])
    _body(doc, "사람이 선택한 설정이 자동 추정보다 우선합니다. 설정이 잘못되면 표시되는 검토 항목도 달라질 수 있습니다.")

    _heading(doc, "3. 조항별 검토 화면 읽는 법", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, text in enumerate(("화면 안내", "뜻")):
        _shade(table.rows[0].cells[i], LIGHT_BLUE)
        _add_text(table.rows[0].cells[i], text, True, DARK_BLUE)
    rows = [
        ("관련 원문 찾음", "체크 항목과 관련된 계약서 위치를 찾았다는 뜻입니다. 문제없음이 확정된 것은 아닙니다."),
        ("원문 확인 필요", "관련 가능성이 있는 원문을 찾았지만 사람이 내용을 확인해야 합니다."),
        ("적용·보완 판단 필요", "이 계약에 적용되는지, 계약 내용을 보완해야 하는지 사람이 판단해야 합니다."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        _add_text(cells[0], left, True)
        _add_text(cells[1], right)
    _table_geometry(table, [2700, 6660])
    _callout(doc, "기억할 점", "시스템은 원문 위치를 안내할 뿐입니다. 중요한 항목은 반드시 계약서 원문을 직접 확인하세요.", GOLD)

    _heading(doc, "항목의 중요도", 2)
    importance_bullets = _new_numbering(doc, "bullet", "•")
    _bullet(doc, "필수: 반드시 확인해야 할 항목", importance_bullets)
    _bullet(doc, "권장: 계약 조건상 확인하는 것이 좋은 항목", importance_bullets)
    _bullet(doc, "참고: 검토에 도움이 되는 참고 항목", importance_bullets)

    _heading(doc, "4. 판정과 코멘트 남기기", 1)
    _heading(doc, "이상없음", 2)
    _body(doc, "계약 내용을 확인한 결과 별도 보완이 필요하지 않은 경우 선택합니다. 필요하면 사유를 선택하고 확인 메모를 남깁니다.")
    ok_bullets = _new_numbering(doc, "bullet", "•")
    _bullet(doc, "반영되어 있음: 계약서 내용으로 충분히 확인한 경우", ok_bullets)
    _bullet(doc, "해당사항 없음: 이 계약에는 해당 항목이 적용되지 않는 경우", ok_bullets)

    _heading(doc, "검토의견", 2)
    _body(doc, "수정, 보완, 협의 또는 추가 확인이 필요한 경우 선택합니다. 이유와 필요한 조치를 코멘트에 적습니다.")
    _callout(doc, "좋은 코멘트", "문제만 적기보다 '왜 문제인지'와 '어떻게 고칠지'를 함께 적으면 담당자와 협의하기 쉽습니다.")
    _body(doc, "판정을 선택해도 항목 위치는 바뀌지 않습니다. 코멘트는 입력 중 자동 저장되며 길이 제한 없이 종합리포트에 표시됩니다.")

    _heading(doc, "5. 종합리포트 확인하기", 1)
    _body(doc, "종합리포트는 시스템의 원문 검색 결과가 아니라 사람의 검토 진행과 판정을 중심으로 보여줍니다.")
    report_bullets = _new_numbering(doc, "bullet", "•")
    _bullet(doc, "검토 완료 상태: 사람이 판정해야 할 항목이 남아 있는지", report_bullets)
    _bullet(doc, "검토자 판정: 이상없음과 검토의견을 몇 건 남겼는지", report_bullets)
    _bullet(doc, "형식 확인 필요: 상호, 대표자, 주소 표기에서 확인할 부분이 있는지", report_bullets)
    _body(doc, "추가 검토 필요 항목이 남아 있으면 [조항별 검토로 이동]을 눌러 판정합니다. 필요한 검토가 모두 끝나면 [검토 마치기]가 활성화됩니다.")
    _body(doc, "상단의 종합 검토의견은 자동 초안입니다. [수정]을 눌러 직접 고칠 수 있으며 긴 코멘트도 잘리지 않습니다.")

    doc.add_page_break()
    _heading(doc, "6. 결과 저장과 공유", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, text in enumerate(("버튼", "용도")):
        _shade(table.rows[0].cells[i], LIGHT_BLUE)
        _add_text(table.rows[0].cells[i], text, True, DARK_BLUE)
    rows = [
        ("검토 마치기", "현재 결과를 브라우저에 저장하고 다음 검토를 위한 지식에 반영합니다."),
        ("검토의견 파일로 내보내기", "팀원이나 담당자에게 전달할 검토 결과 파일을 만듭니다."),
        ("인쇄", "종합리포트를 인쇄하거나 PDF로 저장합니다."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        _add_text(cells[0], left, True)
        _add_text(cells[1], right)
    _table_geometry(table, [3200, 6160])
    _callout(doc, "보관 주의", "중요한 검토 결과는 브라우저 저장에만 의존하지 말고 반드시 파일로 내보내 보관하세요. 다른 PC나 브라우저에는 저장 내용이 자동으로 옮겨지지 않습니다.", GOLD)

    _heading(doc, "7. 변경합의서 검토", 1)
    _body(doc, "변경합의서나 개정합의서를 검토할 때는 입력 화면의 [원계약 첨부]로 원계약도 함께 넣으세요. 원계약에 이미 있는 내용을 전제로 변경된 부분을 검토할 수 있습니다.")

    _heading(doc, "8. 사용 전 마지막 확인", 1)
    checklist = doc.add_table(rows=0, cols=1)
    for text in (
        "계약 유형과 당사 지위가 맞는지 확인했습니까?",
        "필요한 적용 모듈을 선택했습니까?",
        "중요 항목의 계약서 원문을 직접 확인했습니까?",
        "검토의견에 이유와 수정 방향을 적었습니까?",
        "최종 결과를 파일로 내보내 보관했습니까?",
    ):
        cell = checklist.add_row().cells[0]
        _shade(cell, LIGHT_GRAY)
        _add_text(cell, "□  " + text)
    _table_geometry(checklist, [9360])

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


if __name__ == "__main__":
    import sys
    build_guide(sys.argv[1], sys.argv[2])
